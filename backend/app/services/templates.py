from uuid import uuid4

from jinja2 import Template

from app.models import TemplateDefinition, WorkflowTask
from app.services.storage import StorageService, safe_filename


class TemplateRenderer:
    def __init__(self) -> None:
        self.storage = StorageService()

    def render_task(self, task: WorkflowTask, template: TemplateDefinition, output_format: str | None = None) -> str:
        export_format = output_format or template.output_format
        content, content_type = self.render_bytes(task, template, export_format)
        filename = safe_filename(f"{task.id}_{template.id}.{self._extension(export_format)}")
        key = f"exports/{uuid4()}_{filename}"
        self.storage.save_bytes(key, content, content_type)
        return key

    def render_bytes(self, task: WorkflowTask, template: TemplateDefinition, output_format: str | None = None) -> tuple[bytes, str]:
        export_format = (output_format or template.output_format).lower()
        rendered = self.render_text(task, template)
        if export_format == "docx":
            return self._render_docx(rendered), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if export_format == "pdf":
            return self._render_pdf(rendered), "application/pdf"
        if export_format == "json":
            return rendered.encode("utf-8"), "application/json; charset=utf-8"
        if export_format == "markdown":
            return rendered.encode("utf-8"), "text/markdown; charset=utf-8"
        return rendered.encode("utf-8"), "text/plain; charset=utf-8"

    def render_text(self, task: WorkflowTask, template: TemplateDefinition) -> str:
        return Template(template.body).render(
            task=task,
            file=task.file,
            fields=task.extracted_fields or {},
            summary=task.summary or {},
            anomalies=task.anomalies or [],
            tables=task.table_data or [],
        )

    def _render_docx(self, markdown_text: str) -> bytes:
        from io import BytesIO

        from docx import Document

        document = Document()
        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _render_pdf(self, markdown_text: str) -> bytes:
        from io import BytesIO
        from xml.sax.saxutils import escape

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        font_name = "STSong-Light"
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(font_name))
        except Exception:
            font_name = "Helvetica"

        title_style = ParagraphStyle(
            "DocumentTitle",
            fontName=font_name,
            fontSize=18,
            leading=24,
            spaceAfter=8,
        )
        heading_style = ParagraphStyle(
            "DocumentHeading",
            fontName=font_name,
            fontSize=13,
            leading=18,
            spaceBefore=8,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "DocumentBody",
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            spaceAfter=4,
        )

        buffer = BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title="Document Agent Export",
        )
        story = []
        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            if line.startswith("# "):
                story.append(Paragraph(escape(line[2:].strip()), title_style))
            elif line.startswith("## "):
                story.append(Paragraph(escape(line[3:].strip()), heading_style))
            elif line.startswith("### "):
                story.append(Paragraph(escape(line[4:].strip()), heading_style))
            elif line.startswith("- "):
                story.append(Paragraph(f"&#8226; {escape(line[2:].strip())}", body_style))
            else:
                story.append(Paragraph(escape(line), body_style))
        document.build(story)
        return buffer.getvalue()

    def _extension(self, output_format: str) -> str:
        return {"markdown": "md", "txt": "txt", "json": "json", "docx": "docx", "pdf": "pdf"}.get(output_format, "txt")

    def extension(self, output_format: str) -> str:
        return self._extension(output_format.lower())
