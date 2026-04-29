import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.core.config import get_settings
from app.services.ocr import OcrService


@dataclass
class ParsedDocument:
    parser_name: str
    text: str
    tables: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[dict[str, Any]] = field(default_factory=list)


class DocumentParser:
    def parse(self, path: Path, original_filename: str) -> ParsedDocument:
        suffix = Path(original_filename).suffix.lower() or path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
            return self._parse_image(path)
        if suffix in {".docx", ".doc"}:
            return self._parse_docx(path)
        if suffix in {".xlsx", ".xlsm", ".xls"}:
            return self._parse_excel(path)
        if suffix == ".csv":
            return self._parse_csv(path)
        return self._parse_text(path)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        settings = get_settings()
        try:
            from pypdf import PdfReader
        except ImportError:
            fallback = self._parse_text(path)
            fallback.parser_name = "text_fallback"
            fallback.warnings.append({"code": "parser_missing", "message": "pypdf is not installed"})
            return fallback
        reader = PdfReader(str(path))
        pages = []
        warnings = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                warnings.append({"code": "empty_page", "message": f"Page {index} had no extractable text"})
            pages.append(text)
        combined_text = "\n\n".join(pages).strip()
        metadata = {"page_count": len(reader.pages)}
        parser_name = "pypdf"
        if len(combined_text) < settings.ocr_min_text_length:
            ocr = OcrService().ocr_pdf(path)
            warnings.extend(ocr.warnings)
            if ocr.text:
                combined_text = ocr.text
                parser_name = ocr.engine or "ocr"
                metadata = {**metadata, **ocr.metadata, "ocr_applied": True}
            else:
                metadata["ocr_applied"] = False
        return ParsedDocument(
            parser_name=parser_name,
            text=combined_text,
            metadata=metadata,
            warnings=warnings,
        )

    def _parse_image(self, path: Path) -> ParsedDocument:
        ocr = OcrService().ocr_image(path)
        return ParsedDocument(
            parser_name=ocr.engine or "image_ocr",
            text=ocr.text,
            metadata={**ocr.metadata, "ocr_applied": bool(ocr.text)},
            warnings=ocr.warnings,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse Word documents") from exc
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = []
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append({"name": f"table_{table_index}", "rows": rows})
        return ParsedDocument(
            parser_name="python-docx",
            text="\n".join(paragraphs),
            tables=tables,
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables)},
        )

    def _parse_excel(self, path: Path) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to parse Excel workbooks") from exc
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        text_lines = []
        tables = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append(values)
                    text_lines.append(" | ".join(values))
            if rows:
                tables.append({"name": sheet.title, "rows": rows})
        return ParsedDocument(
            parser_name="openpyxl",
            text="\n".join(text_lines),
            tables=tables,
            metadata={"sheet_count": len(workbook.worksheets), "table_count": len(tables)},
        )

    def _parse_csv(self, path: Path) -> ParsedDocument:
        rows = []
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.reader(handle):
                rows.append(row)
        text = "\n".join(" | ".join(cell for cell in row) for row in rows)
        return ParsedDocument(
            parser_name="csv",
            text=text,
            tables=[{"name": "csv", "rows": rows}] if rows else [],
            metadata={"row_count": len(rows)},
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        content = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
            try:
                return ParsedDocument(parser_name=f"text:{encoding}", text=content.decode(encoding))
            except UnicodeDecodeError:
                continue
        return ParsedDocument(parser_name="binary", text="", warnings=[{"code": "unreadable", "message": "File is not text-readable"}])
